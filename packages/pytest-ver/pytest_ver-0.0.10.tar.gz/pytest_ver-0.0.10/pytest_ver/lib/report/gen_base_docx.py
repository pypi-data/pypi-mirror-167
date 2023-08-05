import datetime

import docx
import docx.enum
import docx.enum.section
import docx.enum.style
import docx.enum.table
import docx.enum.text
import docx.shared

from .. import services


# -------------------
## Base class for generating a msword docx
class GenBaseDocx:
    def __init__(self):
        ## holds reference to the current PDF doc
        self._doc = None

    # -------------------
    ## initialize document
    #
    # @return None
    def _doc_init(self):
        self._set_page_layout()
        self._add_styles()
        # TODO how to set headers/footers?
        #    - left, middle, right
        #    - pageno

    # -------------------
    ## set page layout to landscape/portrait
    #
    # @return None
    def _set_page_layout(self):
        section = self._doc.sections[0]
        new_width, new_height = section.page_height, section.page_width
        # TODO make it landscape/portrait as requested
        section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height

        # TODO compare margins to PDF margins?
        section.left_margin = docx.shared.Inches(0.5)
        section.right_margin = docx.shared.Inches(0.5)
        section.top_margin = docx.shared.Inches(1.0)
        section.bottom_margin = docx.shared.Inches(1.0)

    # -------------------
    ## add styles needed for various tables
    #
    # @return None
    def _add_styles(self):
        # Note: use keep_with_next to minimize splitting across pages

        style = self._doc.styles.add_style('ver_table1_header', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = self._doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = docx.shared.Pt(10)
        style.font.bold = True
        style.paragraph_format.space_before = docx.shared.Inches(0.08)
        style.paragraph_format.space_after = docx.shared.Inches(0.08)
        style.paragraph_format.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.SINGLE
        style.paragraph_format.alignment = docx.enum.table.WD_ALIGN_VERTICAL.TOP
        style.paragraph_format.keep_with_next = True

        style = self._doc.styles.add_style('ver_table1_cell', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = self._doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = docx.shared.Pt(10)
        style.font.bold = False
        style.paragraph_format.space_before = docx.shared.Inches(0.0)
        style.paragraph_format.space_after = docx.shared.Inches(0.0)
        style.paragraph_format.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.SINGLE
        style.paragraph_format.alignment = docx.enum.table.WD_ALIGN_VERTICAL.TOP
        style.paragraph_format.keep_with_next = True

        style = self._doc.styles.add_style('ver_table1_desc_cell', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = self._doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = docx.shared.Pt(10)
        style.font.bold = False
        style.paragraph_format.space_before = docx.shared.Inches(0.08)
        style.paragraph_format.space_after = docx.shared.Inches(0.0)
        style.paragraph_format.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.SINGLE
        style.paragraph_format.alignment = docx.enum.table.WD_ALIGN_VERTICAL.TOP
        style.paragraph_format.keep_with_next = True

    # -------------------
    ## generate test run information
    #
    # @return None
    def _gen_test_run_details(self):
        # TODO find BodyText style
        # TODO not blue title?
        self._doc.add_heading('Test Run Details', level=3)

        line = f"{'Test Run Type': <20}: {services.cfg.test_run_type}"
        self._doc.add_paragraph(line, style='List Bullet')

        line = f"{'Test Run ID': <20}: {services.cfg.test_run_id}"
        self._doc.add_paragraph(line, style='List Bullet')

        dts = datetime.datetime.now(datetime.timezone.utc).astimezone().strftime(services.cfg.dts_format)
        line = f"{'Document Generated': <20}: {dts}"
        self._doc.add_paragraph(line, style='List Bullet')

    # -------------------
    ## generate title
    #
    # @param title  the title to draw
    # @return None
    def _gen_title(self, title):
        # TODO how to not make it blue?
        self._doc.add_heading(title, level=3)
